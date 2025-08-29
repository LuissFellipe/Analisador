import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import re
import pandas as pd
from pathlib import Path
from datetime import datetime

try:
    import PyPDF2
except ImportError:
    messagebox.showerror("Erro", "PyPDF2 não instalado. Instale com: pip install PyPDF2")
    raise

try:
    import docx
except ImportError:
    messagebox.showerror("Erro", "python-docx não instalado. Instale com: pip install python-docx")
    raise

def extrair_texto(arquivo_path, progress_callback=None):
    ext = Path(arquivo_path).suffix.lower()
    texto = ""
    try:
        if ext == ".pdf":
            with open(arquivo_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                num_pages = len(reader.pages)
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        texto += page_text + "\n"
                    if progress_callback:
                        progress_callback((i + 1) / num_pages * 100)
        elif ext == ".txt":
            with open(arquivo_path, "r", encoding="utf-8") as f:
                lines = f.readlines()
                total = len(lines)
                for i, line in enumerate(lines):
                    texto += line
                    if progress_callback and total > 0:
                        progress_callback((i + 1) / total * 100)
        elif ext == ".doc":
            doc = docx.Document(arquivo_path)
            total = len(doc.paragraphs)
            for i, para in enumerate(doc.paragraphs):
                texto += para.text + "\n"
                if progress_callback and total > 0:
                    progress_callback((i + 1) / total * 100)
        else:
            messagebox.showerror("Erro", "Tipo de arquivo não suportado.")
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao ler arquivo: {e}")
    return texto

def detectar_dados(texto):
    padroes = {
        "CPF": r"\b\d{3}\.\d{3}\.\d{3}-\d{2}\b",
        "CNPJ": r"\b\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}\b",
        "Email": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
        "Celular": r"(?:\(\d{2}\)\s?)?\d{5}-\d{4}",
        "Cartão de Crédito": r"\b\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4}\b",
        "RG": r"\b\d{2}\.\d{3}\.\d{3}-\d{1}\b",
        "Passaporte": r"\b[A-Z]{2}\d{6}\b"
    }
    resultados = {tipo: {} for tipo in padroes}
    for tipo, regex in padroes.items():
        encontrados = re.findall(regex, texto)
        for valor in encontrados:
            if valor in resultados[tipo]:
                resultados[tipo][valor] += 1
            else:
                resultados[tipo][valor] = 1
    return resultados

def salvar_em_xls(resultados, arquivo_pdf):
    linhas = []
    for tipo, valores in resultados.items():
        for valor, qtd in valores.items():
            linhas.append({"Tipo": tipo, "Valor": valor, "Quantidade": qtd})
    if not linhas:
        messagebox.showinfo("Aviso", "Nenhum dado sensível encontrado para exportar.")
        return
    df = pd.DataFrame(linhas)
    nome_base = Path(arquivo_pdf).stem
    nome_arquivo = f"{nome_base}_dados_sensiveis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    caminho = Path.home() / "Desktop" / nome_arquivo
    try:
        df.to_excel(caminho, index=False)
        messagebox.showinfo("Sucesso", f"Arquivo salvo em:\n{caminho}")
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao salvar arquivo: {e}")

class PDFAnalyzerGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Analisador de Dados Sensíveis em PDF/TXT/DOC")
        self.root.geometry("900x600")
        self.pdf_path = ""
        self.resultados = {}

        frame = ttk.Frame(root, padding=20)
        frame.pack(fill=tk.BOTH,
                   expand=True)

        self.btn_selecionar = ttk.Button(frame, text="Selecionar Arquivo", command=self.selecionar_arquivo)
        self.btn_selecionar.pack(pady=10)

        self.lbl_arquivo = ttk.Label(frame, text="Nenhum arquivo selecionado.")
        self.lbl_arquivo.pack(pady=5)

        self.btn_analisar = ttk.Button(frame, text="Analisar Arquivo", command=self.analisar_arquivo, state=tk.DISABLED)
        self.btn_analisar.pack(pady=10)

        self.progress = ttk.Progressbar(frame, orient="horizontal", length=400, mode="determinate")
        self.progress.pack(pady=10)

        self.tree = ttk.Treeview(frame, columns=("Tipo", "Valor", "Quantidade"), show="headings", height=20)
        self.tree.heading("Tipo", text="Tipo")
        self.tree.heading("Valor", text="Valor Encontrado")
        self.tree.heading("Quantidade", text="Quantidade")
        self.tree.column("Tipo", width=120, anchor="center")
        self.tree.column("Valor", width=400, anchor="center")
        self.tree.column("Quantidade", width=100, anchor="center")
        self.tree.pack(pady=10,
                       fill=tk.BOTH,
                       expand=True)

        self.btn_salvar = ttk.Button(frame, text="Exportar lista para XLS", command=self.salvar_xls, state=tk.DISABLED)
        self.btn_salvar.pack(pady=10)

    def selecionar_arquivo(self):
        arquivo = filedialog.askopenfilename(
            title="Selecione um arquivo",
            filetypes=[
                ("Arquivos suportados", "*.pdf *.txt *.doc"),
                ("PDF files", "*.pdf"),
                ("Text files", "*.txt"),
                ("Word 97-2003", "*.doc")
            ]
        )
        if arquivo:
            self.pdf_path = arquivo
            self.lbl_arquivo.config(text=arquivo)
            self.btn_analisar.config(state=tk.NORMAL)
            self.limpar_resultados()

    def analisar_arquivo(self):
        self.limpar_resultados()
        self.progress["value"] = 0
        self.root.update_idletasks()

        def atualizar_progresso(valor):
            self.progress["value"] = valor
            self.root.update_idletasks()

        texto = extrair_texto(self.pdf_path, progress_callback=atualizar_progresso)
        self.resultados = detectar_dados(texto)
        linhas = []
        for tipo, valores in self.resultados.items():
            for valor, qtd in valores.items():
                linhas.append((tipo, valor, qtd))
        if not linhas:
            messagebox.showinfo("Resultado", "Nenhum dado sensível encontrado.")
            self.btn_salvar.config(state=tk.DISABLED)
            self.progress["value"] = 0
            return

        for linha in linhas:
            self.tree.insert("", tk.END, values=linha)
        self.btn_salvar.config(state=tk.NORMAL)
        self.progress["value"] = 0

    def salvar_xls(self):
        salvar_em_xls(self.resultados, self.pdf_path)

    def limpar_resultados(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
        self.resultados = {}
        self.btn_salvar.config(state=tk.DISABLED)

def main():
    root = tk.Tk()
    app = PDFAnalyzerGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()